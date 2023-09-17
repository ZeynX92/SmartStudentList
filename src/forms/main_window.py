import os.path
import time
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt5 import uic, QtWidgets
from data import db_session
from sqlalchemy import or_
from data.students import Student
from data.documents import Document
from .student import AddStudentForm, EditStudentForm
from .document import AddDocumentForm, EditDocumentForm
from PyQt5.QtWidgets import QMainWindow, QWidget, QTableWidgetItem, QFileDialog
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font
from email_module import send_email


class SmartStudentList(QMainWindow):
    def __init__(self):
        super().__init__()

        uic.loadUi('ui/main.ui', self)

        db_name = "students.db"
        db_session.global_init(f"db/{db_name}")

        self.warn = None
        self.titles = None
        self.titles_1 = None
        self.file_name = None
        self.student_card = None
        self.add_student_form = None
        self.edit_student_form = None

        # Список учеников
        self.tableWidget.setEditTriggers(QtWidgets.QTableWidget.NoEditTriggers)
        self.add_student_button.clicked.connect(self.add_student)
        self.search_student_button.clicked.connect(self.search_student)
        self.tableWidget.cellDoubleClicked.connect(self.edit_student)

        self.search_student()

        # Обработка файлов Импорт Экспорт
        self.lineEdit_4.setReadOnly(True)
        self.lineEdit_2.setReadOnly(True)
        self.lineEdit_5.setReadOnly(True)

        self.list_path = None
        self.table_path = None
        self.table_save_path = None

        self.select_path_button_2.clicked.connect(self.select_save_path_list)
        self.create_list_button.clicked.connect(self.create_list)
        self.create_email_button.clicked.connect(self.create_email)
        self.select_file_button.clicked.connect(self.select_file)
        self.start_cut_button.clicked.connect(self.start_cut)
        self.select_path_button.clicked.connect(self.select_save_path_keys)

        # Хранилище документов
        self.tableWidget_1.setEditTriggers(QtWidgets.QTableWidget.NoEditTriggers)
        self.add_document_button.clicked.connect(self.add_document)
        self.search_button_2.clicked.connect(self.search_document)
        self.tableWidget_1.cellDoubleClicked.connect(self.edit_document)

        self.search_document()

    def add_student(self):
        self.add_student_form = AddStudentForm(self)
        self.add_student_form.show()
        self.setEnabled(False)

    def edit_student(self):
        self.edit_student_form = EditStudentForm(self)
        self.edit_student_form.show()
        self.setEnabled(False)

    def update_table_student(self, result):
        print(f"Here {result}", bool(result))
        self.tableWidget.setSortingEnabled(False)
        if result:
            self.tableWidget.clearContents()
            self.tableWidget.setRowCount(len(result))
            self.tableWidget.setColumnCount(len(result[0]))
            self.titles = ["Фамилия", "Имя", "Отчество", "E-mail", "Телефон", "Заметка"]

            header = self.tableWidget.horizontalHeader()
            self.tableWidget.setHorizontalHeaderLabels(self.titles)
            header.setSectionResizeMode(QtWidgets.QHeaderView.ResizeToContents)

            for i, elem in enumerate(result):
                for j, val in enumerate(elem):
                    self.tableWidget.setItem(i, j, QTableWidgetItem(str(val)))
        else:
            self.warn = Warn("По вашему запросу ничего не найдено!")
            self.warn.show()

        self.tableWidget.resizeColumnsToContents()
        self.tableWidget.setSortingEnabled(True)

    def search_student(self):
        db_sess = db_session.create_session()
        result_raw = db_sess.query(Student).filter(or_(
            Student.surname.contains(self.lineEdit.text()), Student.name.contains(
                self.lineEdit.text()), Student.lastname.contains(self.lineEdit.text()), Student.email.contains(
                self.lineEdit.text()), Student.telephone_number.contains(
                self.lineEdit.text()), Student.note.contains(self.lineEdit.text()))).all()

        result = []
        for student in result_raw:
            result.append([student.surname, student.name, student.lastname, student.email, student.telephone_number,
                           student.note])

        print("Search students: ", sorted(result))
        self.update_table_student(sorted(result))

    def add_document(self):
        self.document_add_form = AddDocumentForm(self)
        self.document_add_form.show()
        self.setEnabled(False)

    def update_table_document(self, result):
        print(f"Here {result}", bool(result))
        self.tableWidget_1.setSortingEnabled(False)
        if result:
            self.tableWidget_1.setRowCount(len(result))
            self.tableWidget_1.setRowCount(len(result))
            self.tableWidget_1.setColumnCount(len(result[0]))
            self.titles_1 = ["Документ", "Описание"]

            header = self.tableWidget.horizontalHeader()
            self.tableWidget_1.setHorizontalHeaderLabels(self.titles_1)
            header.setSectionResizeMode(QtWidgets.QHeaderView.ResizeToContents)

            for i, elem in enumerate(result):
                for j, val in enumerate(elem):
                    self.tableWidget_1.setItem(i, j, QTableWidgetItem(str(val)))
        else:
            self.warn = Warn("По вашему запросу ничего не найдено!")
            self.warn.show()

        self.tableWidget_1.resizeColumnsToContents()
        self.tableWidget_1.setSortingEnabled(True)

    def search_document(self):
        db_sess = db_session.create_session()
        result_raw = db_sess.query(Document).filter(or_(
            Document.title.contains(self.lineEdit_3.text()), Document.description.contains(
                self.lineEdit_3.text()))).all()

        result = []
        for document in result_raw:
            result.append([document.title, document.description])

        self.update_table_document(result)

    def edit_document(self):
        self.edit_document_form = EditDocumentForm(self)
        self.edit_document_form.show()
        self.setEnabled(False)

    def select_file(self):
        self.table_path = QFileDialog.getOpenFileName(self, 'Выбрать файл', '', 'All Files (*);')[0]
        self.lineEdit_5.setText(self.table_path)

    def select_save_path_keys(self):
        self.table_save_path = QFileDialog.getExistingDirectory(self, "Select Directory")
        self.lineEdit_2.setText(self.table_save_path)

    def start_cut(self):
        self.keys_cutted = False
        try:
            table = load_workbook(self.table_path)
            worksheet = table.active
            table_content_raw = tuple(worksheet.rows)
            table_content = [[cell.value for cell in row] for row in table_content_raw]
            print(*table_content, sep='\n')

            for i in range(len(table_content) - 1):
                workbook = Workbook()
                active_worksheet = workbook.active
                active_worksheet.append(table_content[0])
                active_worksheet.append(table_content[i + 1])

                red_font = Font(color='00FF0000', bold=True)

                for cell in active_worksheet["1:1"]:
                    cell.font = red_font

                workbook.save(os.path.join(self.table_save_path,
                                           f'{table_content[i + 1][0].split()[0]} {table_content[i + 1][0].split()[1]} {table_content[i + 1][0].split()[2]}.xlsx'))
                self.label.setText("Аврора успешно нарезала ключи!")
                self.keys_cutted = True

        except Exception:
            self.warn = Warn("Проверьте пути...")
            self.warn.show()

    def create_email(self):
        if self.keys_cutted:
            fails = []
            db_sess = db_session.create_session()
            result_raw = db_sess.query(Student).all()

            for student in result_raw:
                try:
                    send_email(student.email, os.path.join(self.table_save_path,
                                                           f'{student.surname} {student.name} {student.lastname}.xlsx').replace(
                        '\\', '/'))
                    time.sleep(2)
                except Exception:
                    fails.append(f'{student.surname} {student.name} {student.lastname}')
                    continue
            self.label.setText('Данные ученики не смогли получить свои ключи\n' + '\n'.join(
                fails) if fails else "Аврора успешно завершила рассылку")
        else:
            self.warn = Warn("Вы еще не нарезали ключи...")
            self.warn.show()

    def select_save_path_list(self):
        self.list_path = QFileDialog.getExistingDirectory(self, "Select Directory")
        self.lineEdit_4.setText(self.list_path)

    def create_list(self):
        self.label.setText(f"Аврора рада Вас приветствовать!")
        try:
            table_header = None
            doc = docx.Document()

            db_sess = db_session.create_session()
            result_raw = db_sess.query(Student).all()

            result = []
            for student in result_raw:
                if self.comboBox.currentText() == 'Все':
                    table_header = ["ФИO", "E-mail", "Телефон"]
                    result.append(
                        (f'{student.surname} {student.name} {student.lastname}', student.email,
                         student.telephone_number))
                elif self.comboBox.currentText() == 'Только ФИО':
                    table_header = ["Фамилия", "Имя", "Отчество"]
                    result.append((student.surname, student.name, student.lastname))
                elif self.comboBox.currentText() == 'Телефон':
                    table_header = ["ФИO", "Телефон"]
                    result.append((f'{student.surname} {student.name} {student.lastname}', student.telephone_number))
                elif self.comboBox.currentText() == 'E-mail':
                    table_header = ["ФИO", "E-mail"]
                    result.append((f'{student.surname} {student.name} {student.lastname}', student.email))

            items = tuple(sorted(result))

            table = doc.add_table(1, len(items[0]))
            table.style = 'Table Grid'

            head_cells = table.rows[0].cells
            for i, item in enumerate(table_header):
                p = head_cells[i].paragraphs[0]
                p.add_run(item).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for row in items:
                cells = table.add_row().cells
                for i, item in enumerate(row):
                    cells[i].text = str(item)
                    cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cells[i].paragraphs[0].runs[0].font.size = Pt(14)

            doc.save(os.path.join(self.list_path,
                                  f'Список класса ' + \
                                  f'{self.comboBox.currentText() if self.comboBox.currentText() != "Все" else "полный"}' + \
                                  f'.docx'))
            self.label.setText(f"Аврора успешно создала список!")
        except Exception:
            self.warn = Warn("Невозможно перезаписать файл. Проверьте путь и разрешения")
            self.warn.show()


class Warn(QWidget):
    def __init__(self, warn_text: str):
        super().__init__()
        uic.loadUi('ui/warning.ui', self)

        self.label_2.setText(warn_text)
